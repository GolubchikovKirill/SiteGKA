from __future__ import annotations

import datetime
import io
import re
import zipfile
from dataclasses import dataclass

import pytds
import qrcode
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


@dataclass
class QRGeneratorParams:
    server: str
    database: str
    sql_login: str
    sql_password: str
    airport_code: str
    surnames: str | None
    add_login: bool
    both_databases: bool


def _safe_db_name(value: str) -> str:
    if not re.match(r"^[A-Za-z0-9_]+$", value):
        raise ValueError("database name contains unsupported characters")
    return value


def _query_rows(
    *,
    server: str,
    database: str,
    sql_login: str,
    sql_password: str,
    airport_code: str,
    surnames: str | None,
) -> list[dict]:
    safe_db = _safe_db_name(database)
    surname_list = [item.strip() for item in (surnames or "").split(",") if item.strip()]

    query = f"""
        SELECT a.[LOGIN], a.[NAME], a.[PASSWORD], b.[NameExt]
        FROM [{safe_db}].[dbo].[MOL] AS a
        INNER JOIN [{safe_db}].[dbo].[MOLEXT] AS b ON b.[CODE] = a.[CODE]
        WHERE a.[LOGIN] LIKE %s
    """
    params: list[str] = [f"{airport_code}%"]
    if surname_list:
        query += " AND (" + " OR ".join(["a.[NAME] LIKE %s" for _ in surname_list]) + ")"
        params.extend([f"%{item}%" for item in surname_list])
    query += " ORDER BY a.[LOGIN]"

    with pytds.connect(
        server=server,
        database=database,
        user=sql_login,
        password=sql_password,
        autocommit=True,
        as_dict=True,
    ) as conn:
        with conn.cursor() as cur:
            cur.execute(query, tuple(params))
            rows = cur.fetchall() or []
    return list(rows)


def _add_borders_to_cell(cell) -> None:
    tc = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc.append(tc_borders)


def _process_name(value: str) -> str:
    parts = (value or "").strip().split()
    if len(parts) >= 3:
        return " ".join(parts[:2])
    return value or ""


def _generate_qr_png_bytes(data: str) -> bytes:
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(data)
    qr.make(fit=True)
    image = qr.make_image(fill_color="black", back_color="white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _configure_document_layout(doc: Document) -> None:
    """
    Normalize page geometry so Word does not shift table right.

    A fixed A4 portrait with symmetric margins + centered table
    removes the recurring manual correction in desktop Word.
    """
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.gutter = Mm(0)


def _build_word_from_nameext(rows: list[dict], *, airport_code: str, add_login: bool, db_name: str) -> tuple[str, bytes]:
    today = datetime.date.today().strftime("%Y-%m-%d")
    filename = f"Кассиры {db_name}_{airport_code}_{today}.docx"
    doc = Document()
    _configure_document_layout(doc)
    qr_cache: dict[str, bytes] = {}

    for i in range(0, len(rows), 9):
        if i > 0:
            doc.add_page_break()
        table = doc.add_table(rows=3, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            row.height = Inches(2.75)
        for col in table.columns:
            for cell in col.cells:
                # Keep total width well within printable area to avoid drift.
                cell.width = Inches(2.35)
                _add_borders_to_cell(cell)

        for j, row_data in enumerate(rows[i : i + 9]):
            row_idx = j // 3
            col_idx = j % 3
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            qr_data = str(row_data.get("NameExt") or "")
            qr_png = qr_cache.get(qr_data)
            if qr_png is None:
                qr_png = _generate_qr_png_bytes(qr_data)
                qr_cache[qr_data] = qr_png

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(io.BytesIO(qr_png), width=Inches(2))

            name_par = cell.add_paragraph(_process_name(str(row_data.get("NAME") or "")))
            name_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_par.paragraph_format.space_before = Pt(2)
            name_par.paragraph_format.space_after = Pt(2)
            for run in name_par.runs:
                run.font.size = Pt(10)
                run.font.name = "Arial"

            if add_login:
                login_par = cell.add_paragraph(str(row_data.get("LOGIN") or ""))
                login_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                login_par.paragraph_format.space_before = Pt(2)
                login_par.paragraph_format.space_after = Pt(2)
                for run in login_par.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Arial"

    out = io.BytesIO()
    doc.save(out)
    return filename, out.getvalue()


def _build_word_from_login(rows: list[dict], *, airport_code: str, add_login: bool, db_name: str) -> tuple[str, bytes]:
    today = datetime.date.today().strftime("%Y-%m-%d")
    filename = f"СИП {db_name}_{airport_code}_{today}.docx"
    doc = Document()
    _configure_document_layout(doc)
    qr_cache: dict[str, bytes] = {}

    for i in range(0, len(rows), 9):
        if i > 0:
            doc.add_page_break()
        table = doc.add_table(rows=3, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        for row in table.rows:
            row.height = Inches(2.75)
        for col in table.columns:
            for cell in col.cells:
                # Keep total width well within printable area to avoid drift.
                cell.width = Inches(2.35)
                _add_borders_to_cell(cell)

        for j, row_data in enumerate(rows[i : i + 9]):
            row_idx = j // 3
            col_idx = j % 3
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            login = str(row_data.get("LOGIN") or "")
            qr_png = qr_cache.get(login)
            if qr_png is None:
                qr_png = _generate_qr_png_bytes(login)
                qr_cache[login] = qr_png

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(io.BytesIO(qr_png), width=Inches(2))

            login_par = cell.add_paragraph(login)
            login_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            login_par.paragraph_format.space_before = Pt(2)
            login_par.paragraph_format.space_after = Pt(2)
            for run in login_par.runs:
                run.font.size = Pt(10)
                run.font.name = "Arial"

            name_par = cell.add_paragraph(_process_name(str(row_data.get("NAME") or "")))
            name_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_par.paragraph_format.space_before = Pt(2)
            name_par.paragraph_format.space_after = Pt(2)
            for run in name_par.runs:
                run.font.size = Pt(10)
                run.font.name = "Arial"

            if add_login:
                id_par = cell.add_paragraph(login)
                id_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                id_par.paragraph_format.space_before = Pt(2)
                id_par.paragraph_format.space_after = Pt(2)
                for run in id_par.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Arial"

    out = io.BytesIO()
    doc.save(out)
    return filename, out.getvalue()


def generate_qr_docs_zip(params: QRGeneratorParams) -> bytes:
    databases = [(params.server, params.database, "Duty Free" if "KC01" in params.server else "Duty Paid")]
    if params.both_databases:
        databases = [
            ("DC1-SRV-KC01.regstaer.local", "CashDB51", "Duty Free"),
            ("DC1-SRV-KC02.regstaer.local", "CashDB51", "Duty Paid"),
        ]

    generated_docs: list[tuple[str, bytes]] = []
    for server, database, db_label in databases:
        rows = _query_rows(
            server=server,
            database=database,
            sql_login=params.sql_login,
            sql_password=params.sql_password,
            airport_code=params.airport_code,
            surnames=params.surnames,
        )
        if not rows:
            continue
        generated_docs.append(
            _build_word_from_nameext(
                rows,
                airport_code=params.airport_code,
                add_login=params.add_login,
                db_name=db_label,
            )
        )
        generated_docs.append(
            _build_word_from_login(
                rows,
                airport_code=params.airport_code,
                add_login=params.add_login,
                db_name=db_label,
            )
        )

    if not generated_docs:
        raise ValueError("Нет данных по заданным фильтрам.")

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for filename, content in generated_docs:
            archive.writestr(filename, content)
    return zip_buffer.getvalue()
